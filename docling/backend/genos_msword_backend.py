import logging 
import re
from io import BytesIO
from pathlib import Path
from typing import Any, List, Optional, Union
from collections import defaultdict

# PIL 로깅 비활성화 (FpxImagePlugin 오류 메시지 방지)
logging.getLogger('PIL').setLevel(logging.WARNING)

try:
    from wand.image import Image as WandImage
    from wand.exceptions import WandException
    WAND_AVAILABLE = True
except ImportError:
    WAND_AVAILABLE = False

from docling_core.types.doc.base import BoundingBox
from docling_core.types.doc.document import (
    DoclingDocument,
    DocumentOrigin,
    ImageRef,
    NodeItem,
    ProvenanceItem,
    TableCell,
    TableData,
)
from docling_core.types.doc.labels import DocItemLabel, GroupLabel
from docling_core.types.doc.document import Formatting
from docling_core.types.doc import Size
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.table import Table, _Cell
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl
from typing_extensions import override

from docling.backend.abstract_backend import DeclarativeDocumentBackend
from docling.backend.docx.latex.omml import oMath2Latex
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument

_log = logging.getLogger(__name__)


class GenosMsWordDocumentBackend(DeclarativeDocumentBackend):
    @override
    def __init__(
        self, in_doc: "InputDocument", path_or_stream: Union[BytesIO, Path]
    ) -> None:
        super().__init__(in_doc, path_or_stream)
        self.XML_KEY = (
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        )
        self.xml_namespaces = {
            "w": "http://schemas.microsoft.com/office/word/2003/wordml"
        }
        # Word file:
        self.path_or_stream: Union[BytesIO, Path] = path_or_stream
        self.valid: bool = False
        # Initialise the parents for the hierarchy
        self.max_levels: int = 10
        self.level_at_new_list: Optional[int] = None
        self.parents: dict[int, Optional[NodeItem]] = {}
        self.numbered_headers: dict[int, int] = {}
        self.equation_bookends: str = "<eq>{EQ}</eq>"
        # Track processed textbox elements to avoid duplication
        self.processed_textbox_elements: List[int] = []
        # Track content hash of processed paragraphs to avoid duplicate content
        self.processed_paragraph_content: List[str] = []
        # Track seen section texts for header detection (from hwpx_backend)
        self._seen_section_texts: set[str] = set()
        self.processed_table_elements: set[int] = set()

        for i in range(-1, self.max_levels):
            self.parents[i] = None

        self.level = 0
        self.listIter = 0

        self.history: dict[str, Any] = {
            "names": [None],
            "levels": [None],
            "numids": [None],
            "indents": [None],
        }

        self.docx_obj = None
        try:
            if isinstance(self.path_or_stream, BytesIO):
                self.docx_obj = Document(self.path_or_stream)
            elif isinstance(self.path_or_stream, Path):
                self.docx_obj = Document(str(self.path_or_stream))

            self.valid = True
        except Exception as e:
            raise RuntimeError(
                f"GenosMsWordDocumentBackend could not load document with hash {self.document_hash}"
            ) from e

    @override
    def is_valid(self) -> bool:
        return self.valid

    @classmethod
    @override
    def supports_pagination(cls) -> bool:
        return False

    @override
    def unload(self):
        if isinstance(self.path_or_stream, BytesIO):
            self.path_or_stream.close()

        self.path_or_stream = None

    @classmethod
    @override
    def supported_formats(cls) -> set[InputFormat]:
        return {InputFormat.DOCX}

    @override
    def convert(self) -> DoclingDocument:
        """Parses the DOCX into a structured document model.

        Returns:
            The parsed document.
        """

        origin = DocumentOrigin(
            filename=self.file.name or "file",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            binary_hash=self.document_hash,
        )

        doc = DoclingDocument(name=self.file.stem or "file", origin=origin)

        if not self.is_valid():
            raise RuntimeError(
                f"Cannot convert doc with {self.document_hash} because the backend failed to init."
            )

        # --- 1) 각 섹션의 헤더(header.xml) 순회 ---
        for section in self.docx_obj.sections:
            header_el = section.header._element
            doc = self._walk_linear(header_el, self.docx_obj, doc)

        # --- 2) 본문(document.xml) 전체 처리 ---
        body_el = self.docx_obj.element.body
        doc = self._walk_linear(body_el, self.docx_obj, doc)

        # --- 3) 각 섹션의 푸터(footer.xml) 순회 ---
        for section in self.docx_obj.sections:
            footer_el = section.footer._element
            doc = self._walk_linear(footer_el, self.docx_obj, doc)
        doc.pages[1] =  doc.add_page(page_no=1, size=Size(width=595, height=842))
        return doc


    def _update_history(
        self,
        name: str,
        level: Optional[int],
        numid: Optional[int],
        ilevel: Optional[int],
    ):
        self.history["names"].append(name)
        self.history["levels"].append(level)
        self.history["numids"].append(numid)
        self.history["indents"].append(ilevel)

    def _prev_name(self) -> Optional[str]:
        return self.history["names"][-1]

    def _prev_level(self) -> Optional[int]:
        return self.history["levels"][-1]

    def _prev_numid(self) -> Optional[int]:
        return self.history["numids"][-1]

    def _prev_indent(self) -> Optional[int]:
        return self.history["indents"][-1]

    def _get_level(self) -> int:
        """Return the first None index."""
        for k, v in self.parents.items():
            if k >= 0 and v is None:
                return k
        return 0

    def _str_to_int(
        self, s: Optional[str], default: Optional[int] = 0
    ) -> Optional[int]:
        if s is None:
            return None
        try:
            return int(s)
        except ValueError:
            return default

    def _split_text_and_number(self, input_string: str) -> list[str]:
        match = re.match(r"(\D+)(\d+)$|^(\d+)(\D+)", input_string)
        if match:
            parts = list(filter(None, match.groups()))
            return parts
        else:
            return [input_string]

    def _get_numId_and_ilvl(
        self, paragraph: Paragraph
    ) -> tuple[Optional[int], Optional[int]]:
        # Access the XML element of the paragraph
        numPr = paragraph._element.find(
            ".//w:numPr", namespaces=paragraph._element.nsmap
        )

        if numPr is not None:
            # Get the numId element and extract the value
            numId_elem = numPr.find("w:numId", namespaces=paragraph._element.nsmap)
            ilvl_elem = numPr.find("w:ilvl", namespaces=paragraph._element.nsmap)
            numId = numId_elem.get(self.XML_KEY) if numId_elem is not None else None
            ilvl = ilvl_elem.get(self.XML_KEY) if ilvl_elem is not None else None

            return self._str_to_int(numId, None), self._str_to_int(ilvl, None)

        return None, None  # If the paragraph is not part of a list

    def _get_heading_and_level(self, style_label: str) -> tuple[str, Optional[int]]:
        parts = self._split_text_and_number(style_label)

        if len(parts) == 2:
            parts.sort()
            label_str: str = ""
            label_level: Optional[int] = 0
            if parts[0].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[1], None)
            if parts[1].strip().lower() == "heading":
                label_str = "Heading"
                label_level = self._str_to_int(parts[0], None)
            return label_str, label_level

        return style_label, None

    def _get_label_and_level(self, paragraph: Paragraph) -> tuple[str, Optional[int]]:
        if paragraph.style is None:
            return "Normal", None

        label = paragraph.style.style_id
        name = paragraph.style.name
        base_style_label = None
        base_style_name = None
        if base_style := getattr(paragraph.style, "base_style", None):
            base_style_label = base_style.style_id
            base_style_name = base_style.name

        if label is None:
            return "Normal", None

        if ":" in label:
            parts = label.split(":")
            if len(parts) == 2:
                return parts[0], self._str_to_int(parts[1], None)

        if "heading" in label.lower():
            return self._get_heading_and_level(label)
        if "heading" in name.lower():
            return self._get_heading_and_level(name)
        if base_style_label and "heading" in base_style_label.lower():
            return self._get_heading_and_level(base_style_label)
        if base_style_name and "heading" in base_style_name.lower():
            return self._get_heading_and_level(base_style_name)

        return label, None

    @classmethod
    def _get_format_from_run(cls, run: Run) -> Optional[Formatting]:
        # The .bold and .italic properties are booleans, but .underline can be an enum
        # like WD_UNDERLINE.THICK (value 6), so we need to convert it to a boolean
        has_bold = run.bold or False
        has_italic = run.italic or False
        # Convert any non-None underline value to True
        has_underline = bool(run.underline is not None and run.underline)

        return Formatting(
            bold=has_bold,
            italic=has_italic,
            underline=has_underline,
        )

    def _walk_linear(
        self,
        body: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> DoclingDocument:
        for element in body:

            # Check for Inline Images (blip elements)
            namespaces = {
                "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
                "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
                "v": "urn:schemas-microsoft-com:vml",
                "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
                "w10": "urn:schemas-microsoft-com:office:word",
                "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
            }
            xpath_expr = etree.XPath(".//a:blip", namespaces=namespaces)
            drawing_blip = xpath_expr(element)
            # Skip the fallback inside mc:AlternateContent
            tag = etree.QName(element).localname
            if tag == "AlternateContent":
                # find the mc:Choice branch and process only that
                choice = element.find("mc:Choice", namespaces=namespaces)
                if choice is not None:
                    # inline its children into our loop
                    for child in choice:
                        doc = self._walk_linear([child], docx_obj, doc)
                # skip the rest (Fallback)
                continue
            tag_name = etree.QName(element).localname
            # Check for shape content (including textboxes and other shapes)
            # Only process if the element hasn't been processed before
            element_id = id(element)
            if element_id not in self.processed_textbox_elements:
                # Modern Word textboxes
                txbx_xpath = etree.XPath(
                    ".//w:txbxContent|.//v:textbox//w:p", namespaces=namespaces
                )
                textbox_elements = txbx_xpath(element)

                # No modern textboxes found, check for alternate/legacy textbox formats
                if not textbox_elements and tag_name in ["drawing", "pict"]:
                    # Additional checks for textboxes in DrawingML and VML formats
                    alt_txbx_xpath = etree.XPath(
                        ".//wps:txbx//w:p|.//w10:wrap//w:p|.//a:p//a:t",
                        namespaces=namespaces,
                    )
                    textbox_elements = alt_txbx_xpath(element)

                    # Check for shape text that's not in a standard textbox
                    if not textbox_elements:
                        shape_text_xpath = etree.XPath(
                            ".//a:bodyPr/ancestor::*//a:t|.//a:txBody//a:t",
                            namespaces=namespaces,
                        )
                        shape_text_elements = shape_text_xpath(element)
                        if shape_text_elements:
                            # Create custom text elements from shape text
                            text_content = " ".join(
                                [t.text for t in shape_text_elements if t.text]
                            )
                            if text_content.strip():
                                # Create a paragraph-like element to process with standard handler
                                level = self._get_level()
                                shape_group = doc.add_group(
                                    label=GroupLabel.SECTION,
                                    parent=self.parents[level - 1],
                                    name="shape-text",
                                )
                                doc.add_text(
                                    label=DocItemLabel.PARAGRAPH,
                                    parent=shape_group,
                                    text=text_content,
                                    prov=ProvenanceItem(
                                    page_no=1,
                                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                    charspan=(0, 0)
                                )
                                )

                if textbox_elements:
                    # Mark the parent element as processed
                    self.processed_textbox_elements.append(element_id)
                    # Also mark all found textbox elements as processed
                    for tb_element in textbox_elements:
                        self.processed_textbox_elements.append(id(tb_element))

                    self._handle_textbox_content(textbox_elements, docx_obj, doc)

            # Check for shape content (similar to hwpx_backend's _process_rect)
            if tag_name in ["drawing", "pict"] and element_id not in self.processed_textbox_elements:
                self._handle_shape_content(element, docx_obj, doc)

            # Check for Tables - Use enhanced table processing
            if element.tag.endswith("tbl"):
                try:
                    self._handle_tables_enhanced(element, docx_obj, doc)
                except Exception:
                   _log.debug("could not parse a table, broken docx table")

            elif drawing_blip:
                self._handle_pictures(docx_obj, drawing_blip, doc)
                # Check for Text after the Image
                if (tag_name in ["p"]
                    and element.find(".//w:t", namespaces=namespaces) is not None
                ):
                    self._handle_text_elements(element, docx_obj, doc)      
                              
            # Check for the sdt containers, like table of contents
            elif tag_name in ["sdt"]:
                sdt_content = element.find(".//w:sdtContent", namespaces=namespaces)
                if sdt_content is not None:
                    paragraphs = sdt_content.findall(".//w:p", namespaces=namespaces)
                    for p in paragraphs:
                        self._handle_text_elements(p, docx_obj, doc)
            # Check for Text
            elif tag_name in ["p"]:
                # "tcPr", "sectPr"
                self._handle_text_elements(element, docx_obj, doc)
            # else:
            #     _log.debug(f"Ignoring element in DOCX with tag: {tag_name}")
        return doc 

    def _extract_image_from_drawing(
        self, drawing_el: BaseOxmlElement, docx_obj: DocxDocument
    ) -> Optional[ImageRef]:
        """
        <w:drawing> 혹은 VML <v:imagedata> 같은 요소에서
        Word 관계(rId)를 찾아 이미지를 추출합니다.
        """
        # Word ML namespace
        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        }

        # 1) <a:blip> 찾기
        blip = drawing_el.find(".//a:blip", namespaces=ns)
        if blip is None:
            return None

        # 2) 관계 ID 추출
        embed_rId = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not embed_rId or embed_rId not in docx_obj.part.rels:
            return None

        # 3) 이미지 바이너리 가져오기
        image_part = docx_obj.part.rels[embed_rId].target_part
        blob = image_part.blob
        try:
            pil_img = Image.open(BytesIO(blob))
        except UnidentifiedImageError:
            return None

        # 4) ImageRef 생성
        return ImageRef.from_pil(image=pil_img, dpi=72)
    
    def _handle_tables_enhanced(self, element: etree._Element, docx_obj: DocxDocument, doc: DoclingDocument) -> None:
        # 이 element 가 mc:Fallback 계층 안이라면 스킵
        if element.getparent() is not None and etree.QName(element.getparent()).localname == "Fallback":
            return

        # 보수적인 중복 제거: mc:AlternateContent 내부의 명확한 중복만 제거
        parent = element.getparent()
        is_in_alternate_content = False
        
        # Check if this table is inside mc:AlternateContent structure
        ancestor = parent
        while ancestor is not None:
            if etree.QName(ancestor).localname == "AlternateContent":
                is_in_alternate_content = True
                break
            ancestor = ancestor.getparent()
        
        # Only apply content hash duplicate detection for mc:AlternateContent tables
        if is_in_alternate_content:
            try:
                table = Table(element, docx_obj)
                table_content_hash = self._get_table_content_hash(table)
                
                if not hasattr(self, '_processed_table_contents'):
                    self._processed_table_contents = set()
                
                if table_content_hash in self._processed_table_contents:
                    return
                
                self._processed_table_contents.add(table_content_hash)
                
            except Exception as e:
                pass

        eid = id(element)
        #--- 테이블 내부의 텍스트 전체 확인
        # python-docx Table 객체로 변환해서 각 셀을 순회
        table = Table(element, docx_obj)

        # 중복 방지를 위한 ID 체크
        if eid in self.processed_table_elements:
            return
        self.processed_table_elements.add(eid)
                        
        # 1) 기본 테이블 크기
        table = Table(element, docx_obj)
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        ns = element.nsmap

        # 2) table-level detection: 중첩 tbl / 그림이 있는지
        #    - [0]번째는 자기 자신(<w:tbl>)이 잡히기 때문에 [1:]로 실제 중첩 테이블만
        nested_tbls_global = element.findall('.//w:tbl', namespaces=ns)[1:]
        pics_global = (
            element.findall('.//w:drawing', namespaces=ns) +
            element.findall('.//v:imagedata', namespaces=ns)
        )
        table_has_nested = bool(nested_tbls_global)
        table_has_pics   = bool(pics_global)
        
        if num_rows == 1 and num_cols == 1:
            cell_element = table.rows[0].cells[0]
            # In case we have a table of only 1 cell, we consider it furniture
            # And proceed processing the content of the cell as though it's in the document body
            self._walk_linear(cell_element._element, docx_obj, doc)
            return
        # 2) 순수 TableData를 쌓을 객체
        data = TableData(num_rows=num_rows, num_cols=num_cols)

        # 3) 중첩 구조 버퍼: (r, c) → list of (typ, payload)
        cell_buffer = defaultdict(list)
        
        def get_docx_image_bytes(drawing_blip: List[etree._Element]) -> Optional[bytes]:
            # drawing_blip[0] 은 <w:drawing> 엘리먼트
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not rId or rId not in docx_obj.part.rels:
                return None
            return docx_obj.part.rels[rId].target_part.blob
        
        # 4) 각 셀 순회
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                
                # 직속 자식만
                children = list(cell._element.getchildren())
                # 같은 depth에서 중첩 tbl / 그림 검사
                # nested_tbls = [ch for ch in children if etree.QName(ch).localname == "tbl"]
                
                # pics        = cell._element.findall(".//w:drawing", namespaces=ns) \
                #             + cell._element.findall(".//v:imagedata", namespaces=ns) 

                if table_has_nested or table_has_pics:
                    # 중첩 있는 셀: 자식 순서대로 buffer 저장
                    for ch in children:
                        tag = etree.QName(ch).localname
                        if tag == "tbl":
                            cell_buffer[(r_idx, c_idx)].append(("table", ch))
                            continue
                            
                        elif tag == "p":
                            # -- 1) 텍스트 수집
                            texts = [
                                t.text.strip()
                                for t in ch.findall(".//w:t", namespaces=ns)
                                if t.text and t.text.strip()
                            ]
                            if texts:
                                cell_buffer[(r_idx, c_idx)].append(("text", " ".join(texts)))
                                continue

                            # -- 2) drawing 수집
                            drawings = ch.findall(".//w:drawing", namespaces=ns)
                            if drawings:
                                blob = get_docx_image_bytes(drawings)
                                if blob is None:
                                    cell_buffer[(r_idx, c_idx)].append(("picture", None))
                                else:
                                    try:
                                        pil_img = Image.open(BytesIO(blob))
                                        img_ref = ImageRef.from_pil(image=pil_img, dpi=72)
                                        cell_buffer[(r_idx, c_idx)].append(("picture", img_ref))
                                    except UnidentifiedImageError:
                                        # 실패해도 자리 표시
                                        cell_buffer[(r_idx, c_idx)].append(("picture", None)) 
                            continue                                


                    # 만약 버퍼에 뭔가 담겼다면, TableData에 추가하지 않고 continue
                    if (r_idx, c_idx) in cell_buffer:
                        continue
                    
                # 5) 일반 셀: TableData 에 추가
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                data.table_cells.append(
                    TableCell(
                        text=cell_text,
                        row_span=1,
                        col_span=1,
                        start_row_offset_idx=r_idx,
                        end_row_offset_idx=r_idx + 1,
                        start_col_offset_idx=c_idx,
                        end_col_offset_idx=c_idx + 1,
                        column_header=(r_idx == 0),
                        row_header=False,
                    )
                )

        # 6) 버퍼 출력
        parent = self.parents[self._get_level() - 1]
        for (r, c), items in sorted(cell_buffer.items(), key=lambda x: (x[0][0], x[0][1])):
            for typ, payload in items:
                prov = ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=r, r=1, b=r+1),
                    charspan=(0, len(payload) if typ == "text" else 0)
                )
                if typ == "text":
                    doc.add_text(label=DocItemLabel.PARAGRAPH, text=payload, parent=parent, prov=prov)
                    continue
                elif typ == "picture":
                    if payload is None:
                        continue  # 자리 표시가 없으면 스킵
                    else:
                        doc.add_picture(parent=parent, image=payload, caption=None, prov=prov)
                        continue
                elif typ == "table":
                    # 중첩 테이블(가장 안쪽)만 실제 TableData로 재귀 처리
                    self._handle_tables_enhanced(payload, docx_obj, doc)
                    continue

        # 7) TableData 형태로 출력 (중첩 없는 가장 바깥쪽만)
        if data.table_cells:
            doc.add_table(data=data, parent=parent, prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=num_rows),
                charspan=(0, 0)
            ))

    def _should_fallback_table_to_text(self, table: Table, docx_obj: DocxDocument) -> bool:
        """
        Determine if the table should be processed as text instead of a table.
        Returns True if table has complex nested structures that are better processed as text.
        """
        # Check for 1x1 tables
        if len(table.rows) == 1 and len(table.columns) == 1:
            return True
            
        complex_structure_count = 0
        
        for row in table.rows:
            for cell in row.cells:
                # Check for nested tables (currently disabled)
                nested_tables = self._find_nested_tables_in_cell(cell)
                if nested_tables:
                    complex_structure_count += len(nested_tables)
                
                # Check for images (pic elements)
                namespaces = {
                    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
                    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
                }
                
                # Look for picture elements
                xpath_expr = etree.XPath(".//pic:pic", namespaces=namespaces)
                pics = xpath_expr(cell._element)
                if pics:
                    complex_structure_count += len(pics)
                    
                # Look for drawing elements which often contain images
                xpath_expr = etree.XPath(".//w:drawing", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                drawings = xpath_expr(cell._element)
                if drawings:
                    complex_structure_count += len(drawings)
        
        # If we find multiple complex structures, fallback to text
        return complex_structure_count >= 2

    def _pop_table_to_text(self, table: Table, doc: DoclingDocument) -> None:
        """
        Pop current table and process its content as regular text.
        This is the core pop mechanism requested by the user.
        """
        level = self._get_level()
        parent = self.parents[level - 1] if level > 0 else None
        
        all_text_parts = []
        
        for row in table.rows:
            row_text_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                if cell_text:
                    row_text_parts.append(cell_text)
            
            if row_text_parts:
                # Join cell contents with spaces for each row
                row_text = " ".join(row_text_parts)
                all_text_parts.append(row_text)
        
        # Combine all row texts and add as regular text
        if all_text_parts:
            final_text = "\n".join(all_text_parts)
            
            # Check for duplicate content before adding
            if not self._is_duplicate_content(final_text):
                doc.add_text(
                    label=DocItemLabel.TEXT,
                    text=final_text,
                    parent=parent,
                    prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, len(final_text))
                    )
                )
    

    def _process_table_as_text(self, table: Table, doc: DoclingDocument) -> None:
        """
        Process table content as regular text instead of table structure.
        """
        level = self._get_level()
        parent = self.parents[level - 1] if level > 0 else None
        
        all_text_parts = []
        
        for row in table.rows:
            row_text_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                if cell_text:
                    row_text_parts.append(cell_text)
            
            if row_text_parts:
                # Join cell contents with spaces for each row
                row_text = " ".join(row_text_parts)
                all_text_parts.append(row_text)
        
        # Combine all row texts
        if all_text_parts:
            full_text = "\n".join(all_text_parts)
            doc.add_text(
                label=DocItemLabel.TEXT,
                text=full_text,
                parent=parent,
                prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, len(full_text))
                )
            )

    def _find_nested_tables_in_cell(self, cell: _Cell) -> List[BaseOxmlElement]:
        """Find nested tables within a cell. Temporarily disabled to fix pop issue."""
        # TEMPORARY: Disable nested table detection to fix the comprehensive income statement issue
        # Real nested tables are very rare in practice, and the current detection was causing
        # adjacent tables to be incorrectly identified as nested tables
        return []

    def _find_images_in_cell(self, cell: _Cell, docx_obj: DocxDocument) -> List[Any]:
        """Find images within a cell."""
        images = []
        namespaces = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        
        # Look for blip elements (images)
        xpath_expr = etree.XPath(".//a:blip", namespaces=namespaces)
        blips = xpath_expr(cell._element)
        images.extend(blips)
        
        return images

    def _get_table_content_hash(self, table: Table) -> str:
        """
        Generate a hash of table content to detect duplicate tables.
        """
        import hashlib
        
        content_parts = []
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = self._extract_cell_text_with_sdt(cell).strip()
                row_parts.append(cell_text)
            content_parts.append("|".join(row_parts))
        
        table_content = "\n".join(content_parts)
        return hashlib.md5(table_content.encode('utf-8')).hexdigest()

    def _get_text_content_hash(self, text: str) -> str:
        """
        Generate a hash of text content to detect duplicate content.
        """
        import hashlib
        
        # Normalize text: remove extra whitespace and convert to lowercase
        normalized_text = re.sub(r'\s+', ' ', text.strip().lower())
        return hashlib.md5(normalized_text.encode('utf-8')).hexdigest()

    def _is_duplicate_content(self, text: str) -> bool:
        """
        Check if text content is duplicate based on hash.
        """
        if not text or len(text.strip()) < 10:  # Skip very short texts
            return False
            
        text_hash = self._get_text_content_hash(text)
        
        if not hasattr(self, "_processed_text_contents"):
            self._processed_text_contents = set()
            
        if text_hash in self._processed_text_contents:
            return True
            
        self._processed_text_contents.add(text_hash)
        return False

    def _extract_cell_text_with_sdt(self, cell: _Cell) -> str:
        """
        Extract text from cell including content inside SDT (Structured Document Tags).
        This is needed for cells that have form fields or content controls.
        """
        from lxml import etree
        
        namespaces = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        }
        
        # First try the normal cell text
        cell_text = cell.text
        
        # If cell text is empty or very short, look for SDT content
        if not cell_text or len(cell_text.strip()) < 3:
            # Look for all text elements including those inside SDT
            xpath_expr = etree.XPath(".//w:t", namespaces=namespaces)
            text_elements = xpath_expr(cell._element)
            
            all_texts = []
            for t_elem in text_elements:
                if t_elem.text:
                    all_texts.append(t_elem.text)
            
            if all_texts:
                cell_text = "".join(all_texts)
        
        return cell_text if cell_text else "" 

    def _get_paragraph_elements(self, paragraph: Paragraph):
        """
        Extract paragraph elements along with their formatting and hyperlink
        """

        # for now retain empty paragraphs for backwards compatibility:
        if paragraph.text.strip() == "":
            return [("", None, None)]

        paragraph_elements: list[
            tuple[str, Optional[Formatting], Optional[Union[AnyUrl, Path]]]
        ] = []
        group_text = ""
        previous_format = None

        # Iterate over the runs of the paragraph and group them by format
        for c in paragraph.iter_inner_content():
            if isinstance(c, Hyperlink):
                text = c.text
                hyperlink = Path(c.address)
                format = self._get_format_from_run(c.runs[0])
            elif isinstance(c, Run):
                text = c.text
                hyperlink = None
                format = self._get_format_from_run(c)
            else:
                continue

            if (len(text.strip()) and format != previous_format) or (
                hyperlink is not None
            ):
                # If the style changes for a non empty text, add the previous group
                if len(group_text.strip()) > 0:
                    paragraph_elements.append(
                        (group_text.strip(), previous_format, None)
                    )
                group_text = ""

                # If there is a hyperlink, add it immediately
                if hyperlink is not None:
                    paragraph_elements.append((text.strip(), format, hyperlink))
                    text = ""
                else:
                    previous_format = format

            group_text += text

        # Format the last group
        if len(group_text.strip()) > 0:
            paragraph_elements.append((group_text.strip(), format, None))

        return paragraph_elements

    def _get_paragraph_position(self, paragraph_element):
        """Extract vertical position information from paragraph element."""
        # First try to directly get the index from w:p element that has an order-related attribute
        if (
            hasattr(paragraph_element, "getparent")
            and paragraph_element.getparent() is not None
        ):
            parent = paragraph_element.getparent()
            # Get all paragraph siblings
            paragraphs = [
                p for p in parent.getchildren() if etree.QName(p).localname == "p"
            ]
            # Find index of current paragraph within its siblings
            try:
                paragraph_index = paragraphs.index(paragraph_element)
                return paragraph_index  # Use index as position for consistent ordering
            except ValueError:
                pass

        # Look for position hints in element attributes and ancestor elements
        for elem in (*[paragraph_element], *paragraph_element.iterancestors()):
            # Check for direct position attributes
            for attr_name in ["y", "top", "positionY", "y-position", "position"]:
                value = elem.get(attr_name)
                if value:
                    try:
                        # Remove any non-numeric characters (like 'pt', 'px', etc.)
                        clean_value = re.sub(r"[^0-9.]", "", value)
                        if clean_value:
                            return float(clean_value)
                    except (ValueError, TypeError):
                        pass

            # Check for position in transform attribute
            transform = elem.get("transform")
            if transform:
                # Extract translation component from transform matrix
                match = re.search(r"translate\([^,]+,\s*([0-9.]+)", transform)
                if match:
                    try:
                        return float(match.group(1))
                    except ValueError:
                        pass

            # Check for anchors or relative position indicators in Word format
            # 'dist' attributes can indicate relative positioning
            for attr_name in ["distT", "distB", "anchor", "relativeFrom"]:
                if elem.get(attr_name) is not None:
                    return elem.sourceline  # Use the XML source line number as fallback

        # For VML shapes, look for specific attributes
        for ns_uri in paragraph_element.nsmap.values():
            if "vml" in ns_uri:
                # Try to extract position from style attribute
                style = paragraph_element.get("style")
                if style:
                    match = re.search(r"top:([0-9.]+)pt", style)
                    if match:
                        try:
                            return float(match.group(1))
                        except ValueError:
                            pass

        # If no better position indicator found, use XML source line number as proxy for order
        return (
            paragraph_element.sourceline
            if hasattr(paragraph_element, "sourceline")
            else None
        )

    def _collect_textbox_paragraphs(self, textbox_elements):
        """Collect and organize paragraphs from textbox elements."""
        processed_paragraphs = []
        container_paragraphs = {}

        for element in textbox_elements:
            element_id = id(element)
            # Skip if we've already processed this exact element
            if element_id in processed_paragraphs:
                continue

            tag_name = etree.QName(element).localname
            processed_paragraphs.append(element_id)

            # Handle paragraphs directly found (VML textboxes)
            if tag_name == "p":
                # Find the containing textbox or shape element
                container_id = None
                for ancestor in element.iterancestors():
                    if any(ns in ancestor.tag for ns in ["textbox", "shape", "txbx"]):
                        container_id = id(ancestor)
                        break

                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []
                container_paragraphs[container_id].append(
                    (element, self._get_paragraph_position(element))
                )

            # Handle txbxContent elements (Word DrawingML textboxes)
            elif tag_name == "txbxContent":
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )
            else:
                # Try to extract any paragraphs from unknown elements
                paragraphs = element.findall(".//w:p", namespaces=element.nsmap)
                container_id = id(element)
                if container_id not in container_paragraphs:
                    container_paragraphs[container_id] = []

                for p in paragraphs:
                    p_id = id(p)
                    if p_id not in processed_paragraphs:
                        processed_paragraphs.append(p_id)
                        container_paragraphs[container_id].append(
                            (p, self._get_paragraph_position(p))
                        )

        return container_paragraphs

    def _handle_textbox_content(
        self,
        textbox_elements: list,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        level = self._get_level()
        textbox_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="textbox",
        )
        original_parent = self.parents[level]
        self.parents[level] = textbox_group

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        # 1) textbox_elements 중 실제 txbxContent 노드만 뽑기
        txbx_contents = [
            el for el in textbox_elements
            if etree.QName(el).localname == "txbxContent"
        ]

        # 2) 각 txbxContent의 자식 노드를 순서대로 순회
        for content in txbx_contents:
            for child in content.getchildren():
                local = etree.QName(child).localname

                if local == "p":
                    # 단락이면 바로 텍스트 처리
                    self._handle_text_elements(
                        child, docx_obj, doc, is_from_textbox=True
                    )

                elif local == "tbl":
                    # 테이블이면 테이블 처리
                    try:
                        self._handle_tables_enhanced(child, docx_obj, doc)
                    except Exception as e:
                        _log.debug(f"텍스트박스 내 테이블 파싱 실패: {e}")

                # (필요시 다른 태그들: tbl, p 외에도 shape/text 등)

        # 부모 복원
        self.parents[level] = original_parent
        return

    def _handle_shape_content(
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
    ) -> None:
        """Process shape content including tables, text, and images within shapes."""
        namespaces = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
            "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
            "v": "urn:schemas-microsoft-com:vml",
            "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
            "w10": "urn:schemas-microsoft-com:office:word",
            "a14": "http://schemas.microsoft.com/office/drawing/2010/main",
        }

        # Look for text content in shapes
        shape_text_xpath = etree.XPath(
            ".//a:t|.//v:textbox//w:t|.//wps:txbx//w:t", namespaces=namespaces
        )
        text_elements = shape_text_xpath(element)
        
        # Look for tables in shapes
        shape_table_xpath = etree.XPath(
            ".//w:tbl|.//a:tbl", namespaces=namespaces
        )
        table_elements = shape_table_xpath(element)
        
        # Look for paragraphs in shapes
        shape_para_xpath = etree.XPath(
            ".//w:p|.//a:p", namespaces=namespaces
        )
        para_elements = shape_para_xpath(element)

        if not text_elements and not table_elements and not para_elements:
            return

        # Extract all text content from the shape
        all_text_parts = []
        for text_elem in text_elements:
            if text_elem.text:
                all_text_parts.append(text_elem.text.strip())
        
        full_text = " ".join(all_text_parts).strip()
        
        # For longer text or complex content, create a shape group and process content
        level = self._get_level()
        shape_group = doc.add_group(
            label=GroupLabel.SECTION,
            parent=self.parents[level - 1],
            name="shape-content"
        )
        
        # Set this as the current parent temporarily
        original_parent = self.parents[level]
        self.parents[level] = shape_group

        # Process tables within the shape first
        for table_elem in table_elements:
            try:
                self._handle_tables_enhanced(table_elem, docx_obj, doc)
            except Exception as e:
                    _log.debug(f"Could not parse table in shape: {e}")

        # Process paragraphs within the shape
        for para_elem in para_elements:
            try:
                self._handle_text_elements(para_elem, docx_obj, doc)
            except Exception as e:
                    _log.debug(f"Could not parse paragraph in shape: {e}")

        # If we have text but no structured content, add it as plain text
        if full_text and not table_elements and not para_elements:
            # Check for duplicate content before adding
            if not self._is_duplicate_content(full_text):
                doc.add_text(
                    label=DocItemLabel.PARAGRAPH,
                    text=full_text,
                    parent=shape_group,
                    prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, len(full_text))
                    )
                )

        # Restore original parent
        self.parents[level] = original_parent
        return

    def _handle_equations_in_text(self, element, text):
        only_texts = []
        only_equations = []
        texts_and_equations = []
        for subt in element.iter():
            tag_name = etree.QName(subt).localname
            if tag_name == "t" and "math" not in subt.tag:
                if isinstance(subt.text, str):
                    only_texts.append(subt.text)
                    texts_and_equations.append(subt.text)
            elif "oMath" in subt.tag and "oMathPara" not in subt.tag:
                latex_equation = str(oMath2Latex(subt)).strip()
                if len(latex_equation) > 0:
                    only_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )
                    texts_and_equations.append(
                        self.equation_bookends.format(EQ=latex_equation)
                    )

        if len(only_equations) < 1:
            return text, []

        if (
            re.sub(r"\s+", "", "".join(only_texts)).strip()
            != re.sub(r"\s+", "", text).strip()
        ):
            # If we are not able to reconstruct the initial raw text
            # do not try to parse equations and return the original
            return text, []

        # Insert equations into original text
        # This is done to preserve white space structure
        output_text = text[:]
        init_i = 0
        for i_substr, substr in enumerate(texts_and_equations):
            if len(substr) == 0:
                continue

            if substr in output_text[init_i:]:
                init_i += output_text[init_i:].find(substr) + len(substr)
            else:
                if i_substr > 0:
                    output_text = output_text[:init_i] + substr + output_text[init_i:]
                    init_i += len(substr)
                else:
                    output_text = substr + output_text

        return output_text, only_equations

    def _create_or_reuse_parent(
        self,
        *,
        doc: DoclingDocument,
        prev_parent: Optional[NodeItem],
        paragraph_elements: list,
    ) -> Optional[NodeItem]:
        return (
            doc.add_group(label=GroupLabel.INLINE, parent=prev_parent)
            if len(paragraph_elements) > 1
            else prev_parent
        )

    def _handle_text_elements(  # noqa: C901
        self,
        element: BaseOxmlElement,
        docx_obj: DocxDocument,
        doc: DoclingDocument,
        is_from_textbox: bool = False,
    ) -> None:
        paragraph = Paragraph(element, docx_obj)

        # Skip if from a textbox and this exact paragraph content was already processed
        raw_text = paragraph.text
        if is_from_textbox and raw_text:
            # Create a simple hash of content to detect duplicates
            content_hash = f"{len(raw_text)}:{raw_text[:50]}"
            if content_hash in self.processed_paragraph_content:
                return
            self.processed_paragraph_content.append(content_hash)

        text, equations = self._handle_equations_in_text(element=element, text=raw_text)

        if text is None:
            return
        paragraph_elements = self._get_paragraph_elements(paragraph)
        text = text.strip()

        # Common styles for bullet and numbered lists.
        # "List Bullet", "List Number", "List Paragraph"
        # Identify whether list is a numbered list or not
        # is_numbered = "List Bullet" not in paragraph.style.name
        is_numbered = False
        p_style_id, p_level = self._get_label_and_level(paragraph)
        numid, ilevel = self._get_numId_and_ilvl(paragraph)

        if numid == 0:
            numid = None

        # Handle lists
        if (
            numid is not None
            and ilevel is not None
            and p_style_id not in ["Title", "Heading"]
        ):
            self._add_list_item(
                doc=doc,
                numid=numid,
                ilevel=ilevel,
                elements=paragraph_elements,
                is_numbered=is_numbered,
            )
            self._update_history(p_style_id, p_level, numid, ilevel)
            
        
        elif (
            numid is None
            and self._prev_numid() is not None
            and p_style_id not in ["Title", "Heading"]
        ):  # Close list
            if self.level_at_new_list:
                for key in range(len(self.parents)):
                    if key >= self.level_at_new_list:
                        self.parents[key] = None
                self.level = self.level_at_new_list - 1
                self.level_at_new_list = None
            else:
                for key in range(len(self.parents)):
                    self.parents[key] = None
                self.level = 0

        if p_style_id in ["Title"]:
            for key in range(len(self.parents)):
                self.parents[key] = None
            self.parents[0] = doc.add_text(
                parent=None, label=DocItemLabel.TITLE, text=text, prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                charspan=(0, 0)
                )
            )
            
        elif "Heading" in p_style_id:
            style_element = getattr(paragraph.style, "element", None)
            if style_element is not None:
                is_numbered_style = (
                    "<w:numPr>" in style_element.xml or "<w:numPr>" in element.xml
                )
            else:
                is_numbered_style = False
            self._add_header(doc, p_level, text, is_numbered_style)

        elif len(equations) > 0:
            if (raw_text is None or len(raw_text.strip()) == 0) and len(text) > 0:
                # Standalone equation
                level = self._get_level()
                doc.add_text(
                    label=DocItemLabel.FORMULA,
                    parent=self.parents[level - 1],
                    text=text.replace("<eq>", "").replace("</eq>", ""),
                    prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                charspan=(0, 0)
                    )
                )
            else:
                # Inline equation
                level = self._get_level()
                inline_equation = doc.add_group(
                    label=GroupLabel.INLINE, parent=self.parents[level - 1]
                )
                text_tmp = text
                for eq in equations:
                    if len(text_tmp) == 0:
                        break

                    split_text_tmp = text_tmp.split(eq.strip(), maxsplit=1)

                    pre_eq_text = split_text_tmp[0]
                    text_tmp = "" if len(split_text_tmp) == 1 else split_text_tmp[1]

                    if len(pre_eq_text) > 0:
                        doc.add_text(
                            label=DocItemLabel.PARAGRAPH,
                            parent=inline_equation,
                            text=pre_eq_text,
                            prov=ProvenanceItem(
                            page_no=1,
                            bbox=BoundingBox(l=0, t=0, r=1, b=1),
                            charspan=(0, 0)
                            )
                        )
                    doc.add_text(
                        label=DocItemLabel.FORMULA,
                        parent=inline_equation,
                        text=eq.replace("<eq>", "").replace("</eq>", ""),
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )

                if len(text_tmp) > 0:
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=inline_equation,
                        text=text_tmp.strip(),
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )

        elif p_style_id in [
            "Paragraph",
            "Normal",
            "Subtitle",
            "Author",
            "DefaultText",
            "ListParagraph",
            "ListBullet",
            "Quote",
        ]:
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                # Check for duplicate content before adding
                if not self._is_duplicate_content(text):
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )

        else:
            # Text style names can, and will have, not only default values but user values too
            # hence we treat all other labels as pure text
            level = self._get_level()
            parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents.get(level - 1),
                paragraph_elements=paragraph_elements,
            )
            for text, format, hyperlink in paragraph_elements:
                # Check for duplicate content before adding
                if not self._is_duplicate_content(text):
                    doc.add_text(
                        label=DocItemLabel.PARAGRAPH,
                        parent=parent,
                        text=text,
                        formatting=format,
                        hyperlink=hyperlink,
                        prov=ProvenanceItem(
                        page_no=1,
                        bbox=BoundingBox(l=0, t=0, r=1, b=1),
                        charspan=(0, 0)
                        )
                    )

        self._update_history(p_style_id, p_level, numid, ilevel)
        return

    def _add_header(
        self,
        doc: DoclingDocument,
        curr_level: Optional[int],
        text: str,
        is_numbered_style: bool = False,
    ) -> None:
        level = self._get_level()
        if isinstance(curr_level, int):
            if curr_level > level:
                # add invisible group
                for i in range(level, curr_level):
                    self.parents[i] = doc.add_group(
                        parent=self.parents[i - 1],
                        label=GroupLabel.SECTION,
                        name=f"header-{i}",
                    )
            elif curr_level < level:
                # remove the tail
                for key in range(len(self.parents)):
                    if key >= curr_level:
                        self.parents[key] = None

            current_level = curr_level
            parent_level = curr_level - 1
            add_level = curr_level
        else:
            current_level = self.level
            parent_level = self.level - 1
            add_level = 1

        if is_numbered_style:
            if add_level in self.numbered_headers:
                self.numbered_headers[add_level] += 1
            else:
                self.numbered_headers[add_level] = 1
            text = f"{self.numbered_headers[add_level]} {text}"

            # Reset deeper levels
            next_level = add_level + 1
            while next_level in self.numbered_headers:
                self.numbered_headers[next_level] = 0
                next_level += 1

            # Scan upper levels
            previous_level = add_level - 1
            while previous_level in self.numbered_headers:
                # MSWord convention: no empty sublevels
                # I.e., sub-sub section (2.0.1) without a sub-section (2.1)
                # is processed as 2.1.1
                if self.numbered_headers[previous_level] == 0:
                    self.numbered_headers[previous_level] += 1

                text = f"{self.numbered_headers[previous_level]}.{text}"
                previous_level -= 1

        self.parents[current_level] = doc.add_heading(
            parent=self.parents[parent_level],
            text=text,
            level=add_level,
        )
        return

    def _add_list_item(
        self,
        *,
        doc: DoclingDocument,
        numid: int,
        ilevel: int,
        elements: list,
        is_numbered: bool = False,
    ) -> None:
        enum_marker = ""

        level = self._get_level()
        prev_indent = self._prev_indent()
        if self._prev_numid() is None:  # Open new list
            self.level_at_new_list = level

            self.parents[level] = doc.add_group(
                label=GroupLabel.LIST, name="list", parent=self.parents[level - 1]
            )

            # Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )

        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and prev_indent < ilevel
        ):  # Open indented list
            for i in range(
                self.level_at_new_list + prev_indent + 1,
                self.level_at_new_list + ilevel + 1,
            ):
                # Determine if this is an unordered list or an ordered list.
                # Set GroupLabel.ORDERED_LIST when it fits.
                self.listIter = 0
                if is_numbered:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.ORDERED_LIST,
                        name="list",
                        parent=self.parents[i - 1],
                    )
                else:
                    self.parents[i] = doc.add_group(
                        label=GroupLabel.LIST, name="list", parent=self.parents[i - 1]
                    )

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True

            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        elif (
            self._prev_numid() == numid
            and self.level_at_new_list is not None
            and prev_indent is not None
            and ilevel < prev_indent
        ):  # Close list
            for k, v in self.parents.items():
                if k > self.level_at_new_list + ilevel:
                    self.parents[k] = None

            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[self.level_at_new_list + ilevel],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
            self.listIter = 0

        elif self._prev_numid() == numid or prev_indent == ilevel:
            # TODO: Set marker and enumerated arguments if this is an enumeration element.
            self.listIter += 1
            if is_numbered:
                enum_marker = str(self.listIter) + "."
                is_numbered = True
            new_parent = self._create_or_reuse_parent(
                doc=doc,
                prev_parent=self.parents[level - 1],
                paragraph_elements=elements,
            )
            for text, format, hyperlink in elements:
                # Add the list item to the parent group
                doc.add_list_item(
                    marker=enum_marker,
                    enumerated=is_numbered,
                    parent=new_parent,
                    text=text,
                    formatting=format,
                    hyperlink=hyperlink,
                )
        return

    def _handle_pictures(
        self, docx_obj: DocxDocument, drawing_blip: Any, doc: DoclingDocument
    ) -> None:
        def get_docx_image_info(drawing_blip: Any) -> tuple[Optional[bytes], Optional[str]]:
            """이미지 데이터와 형식 정보를 반환합니다."""
            image_data: Optional[bytes] = None
            image_format: Optional[str] = None
            
            rId = drawing_blip[0].get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if rId in docx_obj.part.rels:
                # Access the image part using the relationship ID
                image_part = docx_obj.part.rels[rId].target_part
                image_data = image_part.blob  # Get the binary image data
                # Try to get content type to identify format
                image_format = getattr(image_part, 'content_type', None)
                
            return image_data, image_format

        def is_valid_image_format(image_format: Optional[str], image_data: Optional[bytes]) -> bool:
            """실제 이미지 형식인지 확인합니다."""
            if not image_format or not image_data:
                return False
                
            # XML이나 기타 비이미지 형식 제외
            non_image_formats = [
                'application/xml',
                'text/xml',
                'text/plain',
                'application/json',
                'text/html'
            ]
            
            if image_format.lower() in non_image_formats:
                return False
                
            # 매직 넘버로 실제 이미지인지 확인
            if len(image_data) < 4:
                return False
                
            magic_bytes = image_data[:4]
            
            # XML 시작 패턴 확인
            if magic_bytes.startswith(b'<?xm') or magic_bytes.startswith(b'<xml'):
                return False
                
            # 알려진 이미지 매직 넘버들
            image_signatures = [
                b'\x89PNG',           # PNG
                b'\xff\xd8\xff',      # JPEG
                b'GIF8',              # GIF
                b'BM',                # BMP
                b'RIFF',              # WebP (RIFF 컨테이너)
                b'\x00\x00\x01\x00', # ICO
                b'\xd7\xcd\xc6\x9a', # WMF
                b'\x01\x00\x00\x00', # EMF
                b'II*\x00',          # TIFF (little-endian)
                b'MM\x00*',          # TIFF (big-endian)
            ]
            
            # 매직 넘버 중 하나라도 일치하면 이미지로 간주
            for signature in image_signatures:
                if magic_bytes.startswith(signature):
                    return True
                    
            # content-type이 image로 시작하는 경우에도 시도
            if image_format.lower().startswith('image/'):
                return True
                
            return False

        level = self._get_level()
        # Open the BytesIO object with PIL to create an Image
        image_data, image_format = get_docx_image_info(drawing_blip)
        
        # 이미지 데이터가 없거나 형식이 None인 경우에도 add_picture 호출
        
        # if image_data is None:
        #     doc.add_picture(
        #         parent=self.parents[level - 1],
        #         caption=None,
        #         prov=ProvenanceItem(
        #         page_no=1,
        #         bbox=BoundingBox(l=0, t=0, r=1, b=1),
        #         charspan=(0, 0)
        #             )
        #     )
        #     return
            
        # 실제 이미지인지 확인
        if not is_valid_image_format(image_format, image_data):
            # XML이나 기타 메타데이터는 이미지로 처리하지 않음
            return
            
        try:
            image_bytes = BytesIO(image_data)
            image_bytes.seek(0)  # 포인터를 시작으로 이동
            pil_image = Image.open(image_bytes)
            doc.add_picture(
                parent=self.parents[level - 1],
                image=ImageRef.from_pil(image=pil_image, dpi=72),
                caption=None,
                prov=ProvenanceItem(
                page_no=1,
                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                charspan=(0, 0)
                    )
            )
            
        except (UnidentifiedImageError, OSError) as e:
            print(f"Pillow failed to load image: {e}")
            print(f"Attempting Wand conversion for format: {image_format}")
            
            # WMF/EMF 형식 처리 시도 (Wand 사용)
            if WAND_AVAILABLE and image_format and ('wmf' in image_format.lower() or 'emf' in image_format.lower()):
                try:
                    with WandImage(blob=image_data) as wand_img:
                        # Convert to PNG format
                        wand_img.format = 'png'
                        png_blob = wand_img.make_blob()
                        
                        if png_blob:  # PNG 데이터가 있는지 확인
                            png_bytes = BytesIO(png_blob)
                            png_bytes.seek(0)
                            pil_image = Image.open(png_bytes)
                            doc.add_picture(
                                parent=self.parents[level - 1],
                                image=ImageRef.from_pil(image=pil_image, dpi=72),
                                caption=None,
                                prov=ProvenanceItem(
                                page_no=1,
                                bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                charspan=(0, 0)
                                )
                            )
                            return
                except (WandException, Exception) as wand_error:
                    print(f"Wand conversion failed: {wand_error}")
                
                # 다른 형식도 Wand로 시도
                if WAND_AVAILABLE:
                    try:
                        with WandImage(blob=image_data) as wand_img:
                            # Convert to PNG format
                            wand_img.format = 'png'
                            png_blob = wand_img.make_blob()
                            
                            if png_blob:
                                png_bytes = BytesIO(png_blob)
                                png_bytes.seek(0)
                                pil_image = Image.open(png_bytes)
                                doc.add_picture(
                                    parent=self.parents[level - 1],
                                    image=ImageRef.from_pil(image=pil_image, dpi=72),
                                    caption=None,
                                    prov=ProvenanceItem(
                                    page_no=1,
                                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                                    charspan=(0, 0)
                                    )
                                )
                                return
                    except (WandException, Exception) as wand_error:
                        print(f"Wand fallback conversion failed: {wand_error}")
                
                # 최종적으로 빈 이미지 플레이스홀더 추가
                doc.add_picture(
                    parent=self.parents[level - 1],
                    caption=None,
                    prov=ProvenanceItem(
                    page_no=1,
                    bbox=BoundingBox(l=0, t=0, r=1, b=1),
                    charspan=(0, 0)
                        )
                    )
        return

